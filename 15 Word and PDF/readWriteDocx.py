#! python3
import docx

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

def writeText():
    doc = docx.Document()

    doc.add_heading('Hello, world!', 0)

    paraObj1 = doc.add_paragraph('This is a second paragraph.')
    paraObj2 = doc.add_paragraph('This is a yet another paragraph.')
    paraObj1.add_run(' This text is being added to the second paragraph.')

    doc.paragraphs[1].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

    doc.add_picture('zophie.png', width=docx.shared.Inches(1), height=docx.shared.Cm(4))
    
    doc.save('helloworld.docx')
